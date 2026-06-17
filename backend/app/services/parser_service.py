"""
Parser Service: MIME-dispatched text extraction.

SpAIder's stance is *bring your own parsing* — the canonical path for advanced
users is to pre-parse on their side and POST to ``/ingest`` or
``/ingest/structured``. This service exists only to cover the demo-path
formats (PDF / DOCX / HTML / MD / TXT) with lightweight pure-Python
libraries so the container image stays lean (no PyTorch, no Docling).

Anything beyond these formats — scanned PDFs that need OCR, complex
tables, multi-column academic papers — should be parsed by the caller
with the tool of their choice; SpAIder then ingests the text.

Public surface
--------------
- ``ParseResult(text, hints)`` — dataclass returned by every parser.
- ``parse_sync(payload, content_type=None, filename=None) -> ParseResult`` —
  synchronous entry point.
- ``parse(content, content_type=None, filename=None) -> ParseResult`` —
  *async* entry point retained for backward compatibility with the existing
  ``URLConnector`` / ``UploadConnector`` / ``connector_scheduler`` callers
  that already use ``await parse(...)``. Internally delegates to
  ``parse_sync`` via ``asyncio.to_thread`` so CPU-bound parsing doesn't
  block the event loop.
"""
from __future__ import annotations

import asyncio
import io
import logging
from dataclasses import dataclass, field
from typing import Any, Optional, Union

logger = logging.getLogger(__name__)


@dataclass
class ParseResult:
    """Output of :func:`parse`. ``text`` feeds the graph pipeline; ``hints``
    are non-load-bearing metadata surfaced to the caller (word count, title,
    page count, etc.)."""
    text: str
    hints: dict[str, Any] = field(default_factory=dict)


# ---------------------------------------------------------------------------
# MIME / suffix normalisation
# ---------------------------------------------------------------------------


_MIME_ALIASES: dict[str, str] = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "text/html": "html",
    "application/xhtml+xml": "html",
    "text/markdown": "md",
    "text/x-markdown": "md",
    "text/plain": "txt",
}

_MIME_BY_SUFFIX: dict[str, str] = {
    ".pdf": "pdf",
    ".docx": "docx",
    ".html": "html",
    ".htm": "html",
    ".md": "md",
    ".markdown": "md",
    ".txt": "txt",
}


def resolve_parser(content_type: Optional[str], filename: Optional[str]) -> str:
    """Pick the parser key from MIME (preferred) or file suffix (fallback).

    Unknown text/* MIMEs pass through as plain text rather than raising — a
    client that uploads e.g. ``application/json`` gets its bytes decoded as
    UTF-8, which is almost always what they want for demo flows.
    """
    if content_type:
        base = content_type.split(";", 1)[0].strip().lower()
        if base in _MIME_ALIASES:
            return _MIME_ALIASES[base]
        if base.startswith("text/"):
            return "txt"
    if filename:
        lower = filename.lower()
        for suffix, key in _MIME_BY_SUFFIX.items():
            if lower.endswith(suffix):
                return key
    return "txt"


# ---------------------------------------------------------------------------
# Per-format parsers
# ---------------------------------------------------------------------------


def _parse_pdf(payload: bytes) -> ParseResult:
    """PDF → text via pdfplumber. Tables are rendered as pipe-delimited rows
    (preserves structure without markdown's ambiguity). Scanned / image-only
    PDFs return empty text — that's the signal to the caller to run OCR
    upstream."""
    import pdfplumber

    pages_text: list[str] = []
    table_count = 0
    with pdfplumber.open(io.BytesIO(payload)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            tables = page.extract_tables() or []
            for t in tables:
                table_count += 1
                for row in t:
                    cells = ["" if c is None else str(c).strip() for c in row]
                    page_text += "\n" + " | ".join(cells)
            if page_text.strip():
                pages_text.append(page_text)
        page_count = len(pdf.pages)

    text = "\n\n".join(pages_text).strip()
    return ParseResult(text=text, hints={"page_count": page_count, "table_count": table_count})


def _parse_docx(payload: bytes) -> ParseResult:
    """DOCX → text via python-docx. Paragraphs joined with blank lines;
    tables rendered the same way as the PDF parser for consistency."""
    import docx  # python-docx

    doc = docx.Document(io.BytesIO(payload))
    parts: list[str] = [p.text for p in doc.paragraphs if p.text.strip()]
    table_count = 0
    for table in doc.tables:
        table_count += 1
        for row in table.rows:
            parts.append(" | ".join(cell.text.strip() for cell in row.cells))
    text = "\n\n".join(parts).strip()
    return ParseResult(
        text=text,
        hints={"paragraph_count": len(doc.paragraphs), "table_count": table_count},
    )


# Phrases that commonly appear in JS-only SPA shells — the noscript/fallback
# body that trafilatura lifts when the real content is rendered client-side.
_JS_FALLBACK_MARKERS: tuple[str, ...] = (
    "your browser does not support",
    "please enable javascript",
    "javascript is required",
    "javascript must be enabled",
    "please use a modern browser",
    "does not support es modules",
    "you need to enable javascript to run this app",
)


def _looks_like_js_fallback(text: str) -> bool:
    """True when the extracted text is dominated by JS-bootstrap error copy
    rather than real article content."""
    if not text or len(text) > 400:
        return False
    lowered = text.lower()
    return any(m in lowered for m in _JS_FALLBACK_MARKERS)


def _parse_html(payload: bytes) -> ParseResult:
    """HTML → article text via trafilatura. Aggressively strips nav/ads/
    boilerplate. Detects JS-shell fallback responses and surfaces them
    via ``hints['js_fallback']`` instead of silently ingesting junk."""
    import trafilatura

    raw = payload.decode("utf-8", errors="replace")
    text = trafilatura.extract(raw, include_comments=False, include_tables=True) or ""
    hints: dict[str, Any] = {}
    metadata = trafilatura.extract_metadata(raw)
    if metadata is not None:
        if metadata.title:
            hints["title"] = metadata.title
        if metadata.description:
            hints["description"] = metadata.description
        if metadata.author:
            hints["author"] = metadata.author

    stripped = text.strip()
    if _looks_like_js_fallback(stripped):
        hints["js_fallback"] = True
        hints["js_fallback_preview"] = stripped[:180]
        return ParseResult(text="", hints=hints)
    return ParseResult(text=stripped, hints=hints)


def _parse_markdown(payload: bytes) -> ParseResult:
    """Markdown → plain text. Walks the markdown-it token stream so heading
    structure stays available as a hint without HTML tag noise leaking into
    the graph pipeline."""
    from markdown_it import MarkdownIt

    raw = payload.decode("utf-8", errors="replace")
    md = MarkdownIt("commonmark")
    tokens = md.parse(raw)

    parts: list[str] = []
    heading_count = 0
    for tok in tokens:
        if tok.type == "heading_open":
            heading_count += 1
        if tok.type == "inline" and tok.content:
            parts.append(tok.content)

    text = "\n\n".join(parts).strip()
    word_count = len(text.split())
    return ParseResult(
        text=text, hints={"heading_count": heading_count, "word_count": word_count}
    )


def _parse_passthrough(payload: bytes) -> ParseResult:
    """UTF-8 decode with replacement; no structural hints."""
    text = payload.decode("utf-8", errors="replace").strip()
    return ParseResult(text=text, hints={"word_count": len(text.split())})


# ---------------------------------------------------------------------------
# Public entry points
# ---------------------------------------------------------------------------


_PARSERS = {
    "pdf": _parse_pdf,
    "docx": _parse_docx,
    "html": _parse_html,
    "md": _parse_markdown,
    "txt": _parse_passthrough,
}


class UnsupportedParserError(ValueError):
    """Raised when a payload cannot be dispatched to any parser."""


def parse_sync(
    payload: bytes,
    content_type: Optional[str] = None,
    filename: Optional[str] = None,
) -> ParseResult:
    """Synchronous parse entry point.

    Dispatches on ``content_type`` first, then file suffix, then falls back
    to UTF-8 passthrough. Raises :class:`UnsupportedParserError` only if the
    resolved parser itself fails; unknown types don't raise — they pass
    through as text.
    """
    key = resolve_parser(content_type, filename)
    parser = _PARSERS[key]
    try:
        result = parser(payload)
    except Exception as exc:
        logger.exception("Parser %s failed on %s (%d bytes)", key, filename, len(payload))
        raise UnsupportedParserError(
            f"Failed to parse payload as {key}: {exc}"
        ) from exc
    result.hints.setdefault("parser", key)
    if filename:
        result.hints.setdefault("filename", filename)
    return result


async def parse(
    content: Union[bytes, str],
    content_type: Optional[str] = None,
    filename: Optional[str] = None,
) -> ParseResult:
    """Async parse entry point — preserves the contract used by existing
    callers (``URLConnector``, ``UploadConnector``, ``connector_scheduler``).

    The work is CPU-bound so we offload to a worker thread via
    ``asyncio.to_thread`` instead of blocking the event loop.
    """
    if isinstance(content, str):
        payload = content.encode("utf-8")
    else:
        payload = content
    return await asyncio.to_thread(parse_sync, payload, content_type, filename)


SUPPORTED_MIME_TYPES: tuple[str, ...] = tuple(_MIME_ALIASES.keys())
SUPPORTED_SUFFIXES: tuple[str, ...] = tuple(_MIME_BY_SUFFIX.keys())
